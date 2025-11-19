# -*- coding: utf-8 -*-
# pip install python-docx
import re, difflib
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn

def tokenize(s: str):
    # 공백/비공백 분리로 레이아웃 보존
    return re.findall(r'\s+|[^\s]+', s, flags=re.UNICODE)

def detokenize(tokens):
    return "".join(tokens)

def diff_chunks(original: str, revised: str):
    o, r = tokenize(original), tokenize(revised)
    sm = difflib.SequenceMatcher(a=o, b=r, autojunk=False)
    chunks = []
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == 'equal':
            if j1 != j2:
                chunks.append((detokenize(r[j1:j2]), False))
        elif tag in ('insert', 'replace'):
            if j1 != j2:
                chunks.append((detokenize(r[j1:j2]), True))
        # delete는 수정문 기준 출력에서 제외
    return chunks

def write_docx_with_color(chunks, out_path: str):
    doc = Document()
    # 본문 기본 폰트 설정(한글 대응 폰트; Docs가 없으면 대체)
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.font.name = 'Malgun Gothic'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    style._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')

    para = doc.add_paragraph()
    def new_para():
        return doc.add_paragraph()

    for text, changed in chunks:
        parts = re.split('(\n)', text)
        for part in parts:
            if part == '':
                continue
            if part == '\n':
                para = new_para()
                continue
            run = para.add_run(part)
            if changed:
                run.font.color.rgb = RGBColor(255, 0, 0)  # 변경된 부분 = 빨간색

    doc.save(out_path)
    return out_path

# ===== 여기에 원문/수정본을 붙여넣으세요 =====
email_original = """[대학원 지원 관련 문의] 석사과정 지원 예정자 임소정입니다.

안녕하세요 이준환 교수님, 
2026년 가을학기 서울대학교 언론정보학과 석사과정 진학을 준비 중인 임소정입니다.
저는 미국 University of Iowa에서 영화(Cinema)를 전공하고, SNL Korea와 KBS 시사교양 “더 라이브”에서 조연출로 일한 경력이 있습니다. 

최근 생성형 AI와 VLM 기술이 영상, 영화 제작 전반으로 빠르게 확장되는 흐름을 보며 자연스럽게 Human-AI Interaction 연구에 관심이 생겼고, 교수님 연구들을 접하게 되어 이렇게 연락 드리게 되었습니다.

교수님 연구실의 AI-Human Interaction 관련 연구들을 살펴보면서 전반적으로 많은 흥미를 느꼈지만, 특히 최근에 발표하신 
Journey of Finding the Best Query: Understanding the User Experience of AI Image Generation System
Theatrical Language Processing: Exploring AI-Augmented Improvisational Acting and Scriptwriting with LLMs
두 연구가 개인적으로 크게 인상 깊었습니다.

이미지 생성 과정에서의 프롬프트 탐색 구조나, 즉흥 연기와 스크립트 작성 속에서 AI가 창의적 반응을 이끌어내는 방식은 제가 관심을 가져온 영상 제작, 대본과 대사, 스토리 구성, 영화적 해석 같은 영역과 자연스럽게 연결 되었습니다. 교수님 연구가 최근 이미지·영상 생성 관련 AI-Human Interaction 방향으로도 확장되고 있다는 느낌을 받았고, 특히 Theatrical Language Processing 논문을 보며 영화와 연극 관련 영역에도 관심을 가지고 계신 것 같아 저 역시 더 깊은 관심을 갖고 연락 드리게 되었습니다.

저는 영화 전공 수업들과 방송 조연출로 일하면서 콘티 구성, 샷 디자인, 내러티브 설계, 대본 분석 등 영화·영상 제작의 여러 단계를 경험해왔습니다. 최근 GPT나 Gemini같은 LLM들이 대본과 대사 작업을 지원하고, VLM 모델들이 영상 이해, 요약, 비평까지 다루는 흐름을 보면서 이러한 기술들이 영화적 창작, 해석, 비평 과정에 어떤 영향을 미칠지에 대해 관심이 커졌습니다. Sora, Veo, DALL·E 같은 생성형 모델들은 영상 콘텐츠를 만들어내는 방식 자체를 바꾸고 있다고 느껴지고, 그래서 단순한 생성 능력 뿐만 아니라 영화적 서사, 대본·대사, 평론, 요약 등 영화(Cinema) 전공자가 바라볼 수 있는 다양한 관점에서 AI가 어떻게 창작과 해석을 바꿀 수 있을지 고민하고 있습니다. 저의 이런 생각들이 다소 분산된 것처럼 보일 수 있겠지만, 저에게는 결국 영화라는 매체에서 사람이 AI와 어떤 방식으로 함께 상호작용을 하게 될지에 대한 하나의 큰 질문으로 이어진다고 생각하고 있습니다.

최근에는 Python 기초를 온라인 강의로 배우며 간단한 데이터 분석을 연습하고 있습니다. 부족한 부분이 많지만 연구에 필요한 역량을 꾸준히 준비하고자 노력하고 있습니다.

제가 관심을 갖고 준비하고 있는 연구 방향이 교수님 연구실에서의 연구 주제와 잘 맞을지, 또 제가 앞으로 어떤 부분을 더 준비하면 좋을지 조언을 부탁드리고 싶습니다. 가능하시다면 짧게라도 사전 상담이나 연구 준비에 관해 간단히 여쭙고 싶습니다.

첨부파일에 제 CV를 포함하였습니다.
바쁘신 가운데 읽어주셔서 감사합니다.
좋은 하루 되시길 바랍니다.

임소정 드림
(sojung062400@gmail.com / 010-2123-3621)
"""

email_revised = """[대학원 지원 관련 문의] 2026 가을학기 석사과정 지원 예정자 임소정 드림

이준환 교수님께 안녕하세요.
임소정입니다. 2026년 가을학기 서울대학교 언론정보학과 석사과정 진학을 준비하고 있습니다. University of Iowa에서 영화(Cinema)를 전공했고, SNL Korea와 KBS 시사교양 「더 라이브」에서 조연출로 일했습니다.

관심을 갖게 된 배경을 간단히 말씀드리면 생성형 AI와 VLM이 영상 제작 전반으로 빠르게 확장되는 흐름을 보며 자연스럽게 Human–AI Interaction 분야에 눈을 돌리게 됐습니다. 그러던 중 교수님 연구를 접했고, 특히

Journey of Finding the Best Query: Understanding the User Experience of AI Image Generation System

Theatrical Language Processing: Exploring AI-Augmented Improvisational Acting and Scriptwriting with LLMs
두 작업이 큰 인상을 주었습니다.

제 경험이 교수님 연구와 만나는 지점은 프롬프트 탐색 경험 설계와 즉흥 연기/스크립트에서의 AI 보조가 제가 현장에서 익힌 콘티 구성, 샷 디자인, 내러티브 설계, 대본 분석과 자연스럽게 이어진다는 점입니다. 생성형 모델이 단순한 산출을 넘어 서사·대사·연출의 의사결정에 어떻게 개입하도록 설계될지에 관심이 있습니다. Sora, Veo, DALL·E 등의 진전이 창작–해석–비평의 경계를 바꾸는 과정에서 UX와 협업 인터페이스가 어떤 역할을 해야 하는지도 궁금합니다.

현재 준비 상황은 영화·방송 조연출 경험을 바탕으로 제작 전 과정을 경험했고, 최근에는 Python 기초를 학습하며 간단한 데이터 분석을 연습하고 있습니다. 연구에 필요한 부분을 꾸준히 보완 중입니다.

가능하시다면 조언을 부탁드립니다. 제 관심사가 연구실 주제와 맞는지, 앞으로 어떤 준비가 더 필요할지 여쭙고 싶습니다. 일정이 허락하신다면 짧은 사전 상담(온라인/오프라인) 기회도 부탁드립니다. 편하신 시간대를 알려주시면 맞추겠습니다. CV를 첨부했습니다.

읽어주셔서 감사합니다. 좋은 하루 보내시길 바랍니다.

임소정 드림
sojung062400@gmail.com | 010-2123-3621
"""

if __name__ == "__main__":
    chunks = diff_chunks(email_original, email_revised)
    out = write_docx_with_color(chunks, "email_diff_red.docx")
    print("완료:", out)
